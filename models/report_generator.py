import json
import re
import os
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from docx import Document
import shutil
from pathlib import Path


class ReportDataExtractor:
    """Extrai dados da descrição do recibo para usar no relatório"""

    # Dicionário com nomes dos meses por extenso
    MESES_EXTENSO = {
        1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril",
        5: "maio", 6: "junho", 7: "julho", 8: "agosto",
        9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
    }

    # Dias no mês
    DIAS_MES = {
        1: 31, 2: 28, 3: 31, 4: 30, 5: 31, 6: 30,
        7: 31, 8: 31, 9: 30, 10: 31, 11: 30, 12: 31
    }

    @staticmethod
    def extract_sessions_from_description(descricao: str) -> int:
        """
        Extrai o número de sessões da descrição do recibo.
        Exemplo: "Referente a 04 sessões..." -> 4
                 "Referente a 01 sessão..." -> 1
        """
        try:
            # Procura por "XX sessões" ou "XX sessão"
            match = re.search(r'(\d+)\s+sess(?:ões|ão)',
                              descricao, re.IGNORECASE)
            if match:
                return int(match.group(1))
        except Exception as e:
            print(f"Erro ao extrair número de sessões: {e}")

        return 0

    @staticmethod
    def extract_dates_from_description(descricao: str) -> Tuple[List[str], Optional[str]]:
        """
        Extrai as datas das consultas da descrição.
        Retorna: (lista_de_datas, primeira_data)
        Exemplo: "...nos dias 06/04/2026, 13/04/2026, 20/04/2026 e 27/04/2026."
                 -> (["06/04/2026", "13/04/2026", "20/04/2026", "27/04/2026"], "06/04/2026")
        """
        try:
            # Procura por padrão de datas DD/MM/YYYY
            datas = re.findall(r'(\d{2}/\d{2}/\d{4})', descricao)
            if datas:
                return datas, datas[0]  # Retorna lista e primeira data
        except Exception as e:
            print(f"Erro ao extrair datas: {e}")

        return [], None

    @staticmethod
    def get_last_day_of_month(dia: int, mes: int, ano: int) -> int:
        """
        Retorna o último dia do mês, considerando anos bissextos.
        """
        if mes == 2:
            # Verifica se é ano bissexto
            if ano % 4 == 0 and (ano % 100 != 0 or ano % 400 == 0):
                return 29
            return 28
        return ReportDataExtractor.DIAS_MES.get(mes, 31)

    @staticmethod
    def parse_date(data_str: str) -> Tuple[int, int, int]:
        """
        Parse de data no formato DD/MM/YYYY
        Retorna: (dia, mes, ano)
        """
        try:
            dia, mes, ano = map(int, data_str.split('/'))
            return dia, mes, ano
        except Exception as e:
            print(f"Erro ao fazer parse da data '{data_str}': {e}")
            return 0, 0, 0

    @staticmethod
    def format_cpf(cpf: str) -> str:
        """
        Formata CPF para o padrão "000.000.000-00"
        """
        cpf_limpo = re.sub(r'\D', '', str(cpf))
        if len(cpf_limpo) == 11:
            return f"{cpf_limpo[:3]}.{cpf_limpo[3:6]}.{cpf_limpo[6:9]}-{cpf_limpo[9:]}"
        return cpf_limpo

    @staticmethod
    def extract_report_variables(recibo: Dict, paciente: Dict) -> Dict[str, str]:
        """
        Extrai todas as variáveis necessárias para o relatório.
        """
        variables = {}

        # Determinar qual CPF e nome usar (benef ou pagador)
        cpf_benef = recibo.get('cpf_benef', '').strip()
        cpf_pagador = recibo.get('cpf_pagador', '').strip()

        if cpf_benef and cpf_benef != cpf_pagador:
            # Usar dados de beneficiário (CPF sem formatação: apenas dígitos)
            variables['#CPF'] = re.sub(r"\D", "", str(cpf_benef))
            variables['#NomePac'] = paciente.get('nome_benef', '').strip()
        else:
            # Usar dados de pagador (CPF sem formatação: apenas dígitos)
            variables['#CPF'] = re.sub(r"\D", "", str(cpf_pagador))
            variables['#NomePac'] = paciente.get('nome_pagador', '').strip()

        # Datas de atendimento do paciente
        variables['#DtInicioAtend'] = paciente.get('inicio', '').strip()
        variables['#DtFimAtend'] = paciente.get('fim', '').strip()

        descricao = recibo.get('descricao', '')

        # Número de sessões
        nr_sessoes = ReportDataExtractor.extract_sessions_from_description(
            descricao)
        variables['#NrSessoes'] = str(nr_sessoes)

        # Datas das consultas e primeira data
        datas_consultas, primeira_data = ReportDataExtractor.extract_dates_from_description(
            descricao)

        # Se existir alguma sessão no futuro, usar a data atual no relatório
        hoje = datetime.now().date()
        datas_parsed = []
        for data_str in datas_consultas:
            dia, mes, ano = ReportDataExtractor.parse_date(data_str)
            try:
                data_obj = datetime(ano, mes, dia).date()
            except Exception:
                data_obj = None
            if data_obj is not None:
                datas_parsed.append((data_str, data_obj))

        if any(data_obj > hoje for _, data_obj in datas_parsed):
            hoje_str = hoje.strftime('%d/%m/%Y')
            datas_consultas = [hoje_str]
            primeira_data = hoje_str

        # Formatar datas para exibição (ex: "06/04/2026, 13/04/2026, ...")
        if datas_consultas:
            if len(datas_consultas) == 1:
                variables['#DataDasCons'] = datas_consultas[0]
            else:
                # Formatar com vírgulas e "e" antes do último
                if len(datas_consultas) == 2:
                    variables['#DataDasCons'] = f"{datas_consultas[0]} e {datas_consultas[1]}"
                else:
                    variables['#DataDasCons'] = ", ".join(
                        datas_consultas[:-1]) + f" e {datas_consultas[-1]}"
        else:
            variables['#DataDasCons'] = ''

        # Calcular campos baseados na primeira data
        if primeira_data:
            dia, mes, ano = ReportDataExtractor.parse_date(primeira_data)

            # Último dia do mês
            ult_dia_mes = ReportDataExtractor.get_last_day_of_month(
                dia, mes, ano)
            variables['#UltDiaMesConsultas'] = str(ult_dia_mes)

            # Mês por extenso
            mes_extenso = ReportDataExtractor.MESES_EXTENSO.get(mes, '')
            variables['#MesDasConsultas2'] = mes_extenso

            # Ano
            variables['#AnoDasConsultas2'] = str(ano)

            # Mês numérico com 2 dígitos para nome do arquivo
            variables['#MesNumerico'] = f"{mes:02d}"
        else:
            variables['#UltDiaMesConsultas'] = ''
            variables['#MesDasConsultas2'] = ''
            variables['#AnoDasConsultas2'] = ''
            variables['#MesNumerico'] = ''

        # Forma presencial (por enquanto vazio)
        variables['#FormaPresencial'] = ''

        return variables


class ReportGenerator:
    """Gerador de relatórios a partir de template DOCX"""

    def __init__(self, template_path: str = "modelosRelatorios/RelatorioTemplateCarimboFem.docx"):
        self.template_path = template_path

    def replace_text_in_paragraph(self, paragraph, variables: Dict[str, str]):
        """
        Substitui variáveis dentro de um parágrafo.
        Trata casos onde a variável pode estar dividida entre runs diferentes.
        """
        # Obter todo o texto do parágrafo
        full_text = ''.join([run.text for run in paragraph.runs])

        # Verificar se há variáveis a substituir
        has_replacement = False
        for var_name in variables:
            if var_name in full_text:
                has_replacement = True
                break

        if not has_replacement:
            return

        # Fazer substituições
        for var_name, var_value in variables.items():
            # Substituir no texto completo
            full_text = full_text.replace(var_name, var_value)

        # Limpar runs existentes
        for run in paragraph.runs:
            run.text = ""

        # Adicionar novo texto como um único run
        if paragraph.runs:
            paragraph.runs[0].text = full_text
        else:
            paragraph.add_run(full_text)

    def replace_text_in_document(self, doc, variables: Dict[str, str]):
        """
        Substitui todas as variáveis no documento.
        """
        # Substituir em parágrafos
        for paragraph in doc.paragraphs:
            self.replace_text_in_paragraph(paragraph, variables)

        # Substituir em tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.replace_text_in_paragraph(paragraph, variables)

    def fix_grammar_in_paragraph(self, paragraph):
        """
        Corrige erros gramaticais após substituição de variáveis.
        Considera datas para determinar se usar tempo passado ou futuro:
        - Passado: "foi/foram realizada(s)"
        - Futuro: "será/serão realizada(s)"
        """
        full_text = ''.join([run.text for run in paragraph.runs])
        original_text = full_text

        # Extrair datas do parágrafo (DD/MM/YYYY)
        datas_encontradas = re.findall(r'(\d{2}/\d{2}/\d{4})', full_text)

        # Verificar se alguma data é futura
        from datetime import datetime
        hoje = datetime.now().date()
        tem_data_futura = False

        for data_str in datas_encontradas:
            try:
                dia, mes, ano = map(int, data_str.split('/'))
                data_obj = datetime(ano, mes, dia).date()
                if data_obj > hoje:
                    tem_data_futura = True
                    break
            except Exception:
                pass

        # Extrair número de sessões
        match_sessoes = re.search(
            r'(\d+)\s+sess(?:ões|ão)', full_text, flags=re.IGNORECASE)

        if match_sessoes:
            qtd = match_sessoes.group(1)

            if tem_data_futura:
                # Futuro
                if qtd == '1':
                    full_text = re.sub(
                        r'(?:foi|foram) realizada?s? 1 sess(?:ões|ão)',
                        'será realizada 1 sessão',
                        full_text,
                        flags=re.IGNORECASE
                    )
                else:
                    full_text = re.sub(
                        r'(?:foi|foram) realizada?s? \d+ sess(?:ões|ão)',
                        f'serão realizadas {qtd} sessões',
                        full_text,
                        flags=re.IGNORECASE
                    )
            else:
                # Passado
                if qtd == '1':
                    full_text = re.sub(
                        r'(?:será|serão) realizada?s? 1 sess(?:ões|ão)',
                        'foi realizada 1 sessão',
                        full_text,
                        flags=re.IGNORECASE
                    )
                    full_text = re.sub(
                        r'foram realizadas 1 sess(?:ões|ão)',
                        'foi realizada 1 sessão',
                        full_text,
                        flags=re.IGNORECASE
                    )
                else:
                    full_text = re.sub(
                        r'(?:será|serão) realizada?s? \d+ sess(?:ões|ão)',
                        f'foram realizadas {qtd} sessões',
                        full_text,
                        flags=re.IGNORECASE
                    )

        # Corrigir preposição: "nos dias" <-> "no dia"
        if re.search(r'\b1 sess(?:ões|ão)\b', full_text, flags=re.IGNORECASE):
            full_text = re.sub(r'\bnos dias\b', 'no dia',
                               full_text, flags=re.IGNORECASE)
        else:
            full_text = re.sub(r'\bno dia\b', 'nos dias',
                               full_text, flags=re.IGNORECASE)

        # Se houve mudança, atualizar o parágrafo
        if full_text != original_text:
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = full_text
            else:
                paragraph.add_run(full_text)

    def fix_grammar_in_document(self, doc):
        """
        Corrige erros gramaticais em todo o documento.
        """
        # Corrigir em parágrafos
        for paragraph in doc.paragraphs:
            self.fix_grammar_in_paragraph(paragraph)

        # Corrigir em tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.fix_grammar_in_paragraph(paragraph)

    def generate_report(self, recibo: Dict, paciente: Dict,
                        output_path: str = None) -> bool:
        """
        Gera o relatório substitituindo variáveis do template.
        Salva em formato DOCX primeiro, depois converte para PDF.

        Args:
            recibo: Dict com dados do recibo
            paciente: Dict com dados do paciente
            output_path: Caminho para salvar o PDF (opcional)

        Returns:
            True se bem-sucedido, False caso contrário
        """
        try:
            # Extrair variáveis
            variables = ReportDataExtractor.extract_report_variables(
                recibo, paciente)

            # Carregar template
            if not os.path.exists(self.template_path):
                print(f"Erro: Template não encontrado em {self.template_path}")
                return False

            doc = Document(self.template_path)

            # Substituir variáveis
            self.replace_text_in_document(doc, variables)

            # Corrigir erros gramaticais
            self.fix_grammar_in_document(doc)

            # Definir caminho de saída
            if output_path is None:
                # Gerar nome automaticamente
                nome_beneficiario = variables['#NomePac']
                ano = variables['#AnoDasConsultas2']
                mes = variables['#MesDasConsultas2']

                # Limpar nome para usar como filename
                nome_limpo = re.sub(r'[<>:"/\\|?*]', '',
                                    nome_beneficiario).strip()

                output_path = f"{nome_limpo} - Relatorio {ano}{mes}.docx"

            # Salvar DOCX
            docx_path = output_path if output_path.endswith(
                '.docx') else output_path.replace('.pdf', '.docx')
            doc.save(docx_path)
            print(f"Relatório DOCX gerado: {docx_path}")

            # Converter para PDF
            pdf_path = output_path if output_path.endswith(
                '.pdf') else output_path.replace('.docx', '.pdf')
            success = self.convert_docx_to_pdf(docx_path, pdf_path)

            if success:
                print(f"Relatorio PDF gerado: {pdf_path}")
                # Remover arquivo DOCX intermediário (opcional)
                # os.remove(docx_path)

            return success

        except Exception as e:
            print(f"Erro ao gerar relatorio: {e}")
            return False

    @staticmethod
    def convert_docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
        """
        Converte arquivo DOCX para PDF.
        - Windows: usa Microsoft Word (docx2pdf)
        - Linux/Mac: usa LibreOffice em background
        """
        try:
            import subprocess
            import platform

            if platform.system() == "Windows":
                # No Windows, docx2pdf utiliza o Microsoft Word instalado.
                return ReportGenerator.convert_with_docx2pdf(docx_path, pdf_path)

            # Linux/Mac: conversão com LibreOffice
            cmd = [
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", os.path.dirname(pdf_path) or ".",
                docx_path
            ]

            # Executar comando
            result = subprocess.run(cmd, capture_output=True, timeout=30)

            if result.returncode == 0:
                # Verificar se o arquivo PDF foi criado
                docx_dir = os.path.dirname(docx_path) or "."
                expected_pdf = os.path.join(docx_dir, os.path.basename(
                    docx_path).replace('.docx', '.pdf'))

                if os.path.exists(expected_pdf):
                    # Se o caminho de saída é diferente, mover arquivo
                    if expected_pdf != pdf_path:
                        shutil.move(expected_pdf, pdf_path)
                    return True

            return False

        except Exception as e:
            print(f"Erro ao converter com LibreOffice: {e}")
            return False

    @staticmethod
    def convert_with_docx2pdf(docx_path: str, pdf_path: str) -> bool:
        """
        Converte usando a biblioteca docx2pdf.
        No Windows, isso requer Microsoft Word instalado.
        """
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            return os.path.exists(pdf_path)
        except Exception as e:
            print(f"Erro ao converter com Microsoft Word (docx2pdf): {e}")
            return False


class RecibosReportManager:
    """Gerenciador de relatórios para recibos"""

    def __init__(self):
        self.recibos_storage = None
        self.pacientes_storage = None
        self.report_generator = ReportGenerator()

    def set_storage(self, recibos_storage, pacientes_storage):
        """Define os storages para acessar dados"""
        self.recibos_storage = recibos_storage
        self.pacientes_storage = pacientes_storage

    def get_paciente_by_cpf(self, cpf: str, profissional_cpf: str) -> Optional[Dict]:
        """
        Obtém dados do paciente baseado no CPF e no CPF do profissional.
        """
        try:
            # Construir nome do arquivo do profissional
            if not self.pacientes_storage:
                return None

            # Tentar carregar do arquivo do profissional
            # Primeiro precisamos saber qual é o profissional logado
            # Por enquanto, vamos verificar o arquivo de pacientes padrão
            return None  # Será implementado quando integrar com auth

        except Exception as e:
            print(f"Erro ao buscar paciente: {e}")
            return None

    def generate_report_for_recibo(self, recibo_id: int,
                                   profissional_cpf: str = None,
                                   output_path: str = None) -> bool:
        """
        Gera relatório para um recibo específico.
        """
        try:
            if not self.recibos_storage:
                print("Erro: Storage de recibos não configurado")
                return False

            # Buscar recibo
            recibo = self.recibos_storage.get_recibo(recibo_id)
            if not recibo:
                print(f"Erro: Recibo {recibo_id} não encontrado")
                return False

            # Buscar paciente (será usado com o profissional logado)
            cpf_benef = recibo.get('cpf_benef', '').strip()
            cpf_pagador = recibo.get('cpf_pagador', '').strip()
            cpf_procurado = cpf_benef if cpf_benef and cpf_benef != cpf_pagador else cpf_pagador

            # Para agora, vamos usar um arquivo JSON genérico de pacientes
            # Será melhorado quando integrar com sistema de profissionais
            paciente = self._find_paciente_in_files(cpf_procurado)

            if not paciente:
                print(f"Erro: Paciente com CPF {cpf_procurado} não encontrado")
                return False

            # Gerar relatório
            return self.report_generator.generate_report(recibo, paciente, output_path)

        except Exception as e:
            print(f"Erro ao gerar relatorio para recibo: {e}")
            return False

    def _find_paciente_in_files(self, cpf: str) -> Optional[Dict]:
        """
        Procura um paciente por CPF nos arquivos JSON de profissionais.
        """
        try:
            # Procurar em arquivos de profissionais (formato: {profissional}_pacientes.json)
            for file in os.listdir("."):
                if file.endswith("_pacientes.json"):
                    try:
                        with open(file, 'r', encoding='utf-8') as f:
                            data = json.load(f)
                            pacientes = data.get('pacientes', [])

                            for paciente in pacientes:
                                cpf_benef = paciente.get(
                                    'cpf_benef', '').strip()
                                cpf_pagador = paciente.get(
                                    'cpf_pagador', '').strip()

                                if cpf_benef == cpf or cpf_pagador == cpf:
                                    return paciente
                    except Exception as e:
                        continue

            return None

        except Exception as e:
            print(f"Erro ao procurar paciente: {e}")
            return None
